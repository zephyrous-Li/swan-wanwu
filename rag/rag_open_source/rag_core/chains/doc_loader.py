import os
from typing import List, Optional
from langchain.docstore.document import Document
# from langchain.document_loaders import TextLoader
# from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders import UnstructuredFileLoader
import copy
import docx
from docx import Document as docx_Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import qn
from PIL import Image
import base64
from io import BytesIO
import docx2txt
from utils.minio_utils import upload_local_file
import re
import tiktoken
from langchain_text_splitters import CharacterTextSplitter
# from pymilvus.model.sparse.bm25.tokenizers import build_default_analyzer
import tempfile
from settings import IMAGE_MINIMUM_WIDTH, IMAGE_MINIMUM_HEIGHT
# analyzer = build_default_analyzer(language="zh")

tiktoken_cache_dir = "/opt/tiktoken_cache"
os.environ["TIKTOKEN_CACHE_DIR"] = tiktoken_cache_dir
encoding = tiktoken.encoding_for_model("gpt-4")


def len_fun(text):
    """计算文本长度"""
    return len(encoding.encode(text))


class DOCLoader(TextLoader):
    def load(self) -> List[Document]:
        text = ""
        try:
            text = os.popen("catdoc -s='utf-8' -d='utf-8' " + self.file_path).read()
        except Exception as e:
            raise RuntimeError(f"Error loading {self.file_path}") from e

        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]


class DOCXLoader(TextLoader):
    def load(self) -> List[Document]:
        text = ""
        try:
            # # ================== UnstructuredFileLoader ====================
            # loader = UnstructuredFileLoader(self.file_path, mode="elements")
            # docs = loader.load()
            # text_list=[i.page_content for i in docs]
            # text = "".join(text_list)
            # # ================== UnstructuredFileLoader ====================

            # ================== docx ====================
            text = docx_to_markdown(self.file_path)
            # ================== docx ====================
        except Exception as e:
            raise RuntimeError(f"Error loading {self.file_path}") from e

        metadata = {"source": self.file_path}
        return [Document(page_content=text, metadata=metadata)]

    def custom_load_and_split_doc(self) -> (List[dict]):
        """
        将docx文件加载并拆分为多个段落,分成 chunks 和 sub_chunks 返回
        """
        chunks = []
        sub_chunk = []
        chunk_size = 800
        chunk_overlap = 150
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len_fun,  # 以tokens计算长度
            is_separator_regex=False,
        )
        try:
            # ================== docx custom load and split====================
            text = docx_to_markdown(self.file_path)
            docs = text_splitter.create_documents([text])

            chunk_order = 0
            for doc in docs:  # 添加好到 chunks 和 sub_chunks
                embedding_chunks = []
                chunk_order += 1
                meta_data = {'file_name': self.file_path.split("/")[-1],
                             'chunk_len': len_fun(doc.page_content),
                             'chunk_total_num': len(docs), 'chunk_current_num': chunk_order}

                # sub_chunk.append({'content': doc.page_content, 'embedding_content': doc.page_content,
                #                   'meta_data': meta_data})
                for paragraph in doc.page_content.split('\n\n'):  # 拆分到 sub_chunks
                    embedding_chunks.append(paragraph)
                    # sub_chunk.append({'content': doc.page_content, 'embedding_content': paragraph,
                    #                   'meta_data': meta_data})
                chunks.append({'type': 'text', 'text': doc.page_content, 'meta_data': meta_data, 'embedding_chunks': embedding_chunks})
            # ================== docx custom load and split====================
        except Exception as e:
            raise RuntimeError(f"Error loading docx {self.file_path}") from e

        # 返回 chunks 和 sub_chunks
        return chunks


def docx_to_markdown(docx_path):
    def _table_to_markdown(table):
        markdown = []
        # calculate the total number of columns
        total_cols = max(len(row.cells) for row in table.rows)
        header_row = table.rows[0]
        headers = _parse_row(header_row, total_cols)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        # 收集所有行，用于后续去重
        all_rows = []
        for row in table.rows[1:]:
            row_cells = _parse_row(row, total_cols)
            all_rows.append(row_cells)

        # 对完全相同的行进行去重
        seen_rows = []
        for row_cells in all_rows:
            row_str = "| " + " | ".join(row_cells) + " |"
            if row_str not in seen_rows:
                seen_rows.append(row_str)
                markdown.append(row_str)
        return "\n".join(markdown)

    def _parse_row(row, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        processed_cell_ids = set()  # 用于跟踪已经处理过的单元格对象
        col_index = 0

        for cell in row.cells:
            # 使用 id() 来唯一标识每个单元格对象，因为 python-docx 对合并单元格会返回同一个对象
            cell_id = id(cell)

            # 跳过已经处理过的单元格对象
            if cell_id in processed_cell_ids:
                continue

            processed_cell_ids.add(cell_id)

            # 找到下一个空列的位置
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1

            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break

            cell_content = _parse_cell(cell).strip()
            cell_colspan = cell.grid_span or 1
            cell_colspan = min(cell_colspan, total_cols - col_index)  # 确保不超过总列数

            # 填充单元格内容到相应列
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""

            col_index += cell_colspan

        return row_cells

    def _parse_cell(cell):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = _parse_cell_paragraph(paragraph)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(paragraph):
        paragraph_content = []
        for run in paragraph.runs:
            paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    #  ======== 加载文档 ========
    try:
        doc = docx_Document(docx_path)
    except Exception as e:
        print(f"docx_Document have err:{e}")
        # 提取文本
        text = docx2txt.process(docx_path)
        return text

    part = doc.part

    # Markdown字符串
    markdown_text = ""

    # 遍历文档中的所有段落和表格
    for element in doc.element.body.iterchildren():
        if isinstance(element, docx.oxml.text.paragraph.CT_P):
            p = Paragraph(element, doc)
            # 将段落转换为Markdown格式
            paragraph_text = p.text.strip()

            if paragraph_text:  # 忽略空段落
                markdown_text += f"{paragraph_text}  \n\n"
            if '<w:drawing>' in element.xml:  # 判断是否是插图
                # 正则表达式
                pattern = r'<a:blip\s+r:embed="([^"]+)"'
                # 使用正则表达式搜索
                match = re.search(pattern, element.xml)
                if match:
                    try:
                        image_rid = match.group(1)  # 获取捕获组中的内容
                        print(f"Found r:embed value: {image_rid}")
                        # ========= 将图片转写成 base64 =========
                        image_part = part.related_parts[image_rid]
                        image_bytes = image_part.blob
                        with Image.open(BytesIO(image_bytes)) as img:
                            # 过滤掉小图标
                            width, height = img.size
                            if width < IMAGE_MINIMUM_WIDTH or height < IMAGE_MINIMUM_HEIGHT:  # 设置你认为的小图标大小阈值
                                continue  # 跳过小图标

                            # 创建一个临时文件路径
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                                img_tmp_path = tmp_file.name
                                # 保存图像到临时文件
                                img.save(img_tmp_path)
                                # 上传临时文件
                                minio_result = upload_local_file(img_tmp_path)
                                if minio_result['code'] == 0 and minio_result['download_link']:
                                    image_download_link = minio_result['download_link']
                                    print("====>image_download_link=%s" % image_download_link)
                                    img_url_str = f"![image]({image_download_link})"
                                    markdown_text += f"{img_url_str}"
                                else:
                                    print(f"====>geu image_download_link err:{minio_result}")
                                # img_base64 = base64.b64encode(image_bytes).decode('utf-8')
                                # 删除临时文件
                                os.remove(img_tmp_path)
                    except Exception as e:
                        print(f"====>geu image_download_link err:{e}")
                else:
                    pass
        elif isinstance(element, docx.oxml.table.CT_Tbl):
            # 处理表格
            markdown_text += '\n'  # 在表格前添加空行
            block = Table(element, doc)
            # for row in block.rows:
            #     row_text = '|'.join([cell.text.strip() for cell in row.cells])
            #     markdown_text += f"|{row_text}|\n"
            # markdown_text += '\n'  # 在表格后添加空行
            table_text = _table_to_markdown(block)
            markdown_text += table_text
            markdown_text += '\n'  # 在表格后添加空行
        else:
            pass
    # 打印Markdown文本
    if markdown_text:
        return markdown_text
    else:  # 如果第一种方式没有提取到文本，则使用第二种方式
        print(f"docx_Document markdown_text is None ")
        # 提取文本
        text = docx2txt.process(docx_path)
        return text



if __name__ == "__main__":
    filepath = "./your_file.docx"
    # # 根据文件类型选择加载器
    # if filepath.endswith(".doc"):
    #     loader = DOCLoader(filepath)
    # elif filepath.endswith(".docx"):
    #     loader = DOCXLoader(filepath)
    # docs = loader.load()
    # for doc in docs:
    #     print(doc)

    # ============ 测试使用 docx 读取 文档并转换为 markdown 格式 =============
    docx_to_markdown(filepath)
    print(1)
